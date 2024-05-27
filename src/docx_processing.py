from io import BytesIO
from pathlib import Path
from docx import Document
from docx.enum.text import WD_BREAK

from translation import detect_language, translate_text

class MultipleLanguagesError(Exception):
    """
    Exception raised when multiples languages are detected in the text
    """
    pass

class DocxTranslator:
    """
    Translate a .docx file
    
    Attributes:
        doc (docx.Document): The docx file to translate
        translated_doc (docx.Document): The translated docx file
        detected_language (str): The detected language of the text
        target_language (str): The language code of the target language
        images (list): The list of images in the docx file

    Public methods:
        process_document(self): Process the document
    """
    def __init__(self, file, target_language):
        """
        Init the translator
        """
        self.doc = Document(file)
        self.translated_doc = Document()
        self.detected_language = None
        self.target_language = target_language

        self.images = []

    # PUBLIC METHODS
    def process_document(self):
        """
        Process the document
        """
        for rel in self.doc.part.rels.values():
            if "image" in rel.target_ref:
                image = rel.target_part.blob
                self.images.append(image)
                
        for para in self.doc.paragraphs:
            self.__process_paragraph(para)

        # TODO : add tables during the paragraph processing
        self.translated_doc.add_page_break()
        self.translated_doc.add_heading("TABLES", level=1)
        for i, table in enumerate(self.doc.tables, start=1):
            self.translated_doc.add_heading(f"Table {i}", level=2)
            self.__process_table(table)
            self.translated_doc.add_paragraph("\n\n")

        return self.translated_doc
    
    # PRIVATES METHODS
    def __detect_language(self, text_to_translate):
        """
        Detect the language of the text
        """
        if not text_to_translate:
            return False
        try: 
            self.detected_language = detect_language(text_to_translate)
            return True
        except TypeError as e:
            if str(e) == "sequence item 0: expected str instance, NoneType found":
                return False

    @staticmethod
    def __add_image(paragraph, image_blob):
        """
        Add an image to a paragraph

        Args:
            paragraph (docx.paragraph.Paragraph): The paragraph to add the image to
            image_blob (bytes): The image blob
        """
        run = paragraph.add_run()
        image_stream = BytesIO(image_blob)
        run.add_picture(image_stream)
    
    def __process_paragraph(self, para):
        """
        Process a paragraph

        Args:
            para (docx.paragraph.Paragraph): The paragraph to process
        """
        if "<a:graphicData" in para._p.xml:
            image = self.images.pop(0)
            new_para = self.translated_doc.add_paragraph()
            self.__add_image(new_para, image)

        text_to_translate = para.text
        if self.__detect_language(text_to_translate):
            translated_text = translate_text(text_to_translate, self.detected_language, self.target_language)
        
            if para.style.name.startswith("Heading"):
                new_para = self.translated_doc.add_heading(level=int(para.style.name.split(' ')[-1]))
            else:
                new_para = self.translated_doc.add_paragraph()
            new_para.style = para.style
            
            new_run = new_para.add_run(translated_text)
            for run in para.runs:
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.color.rgb = run.font.color.rgb

                if run.text == '\f':  # Page break character
                    new_run.add_break(WD_BREAK.PAGE)

    def __process_table(self, table):
        """
        Process a table

        Args:
            table (docx.table.Table): The table to process
        """
        new_table = self.translated_doc.add_table(len(table.rows), len(table.columns), style=table.style)

        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                
                new_cell = new_table.cell(row_idx, cell_idx)
                if self.__detect_language(cell.text):
                    translated_text = translate_text(cell.text, self.detected_language, self.target_language)
                    new_cell.text = translated_text

                new_cell.paragraphs[0].style = cell.paragraphs[0].style


if __name__ == "__main__":
    base_dir = Path(__file__).resolve().parent.parent
    samples_dir = base_dir / "samples"
    input_file = samples_dir / "test.docx"
    output_dir = samples_dir / "translated_files"
    output_file = output_dir / "test_fr.docx"
    output_dir.mkdir(parents=True, exist_ok=True)
    with open(input_file, "rb") as f:
        docx_translator = DocxTranslator(f, "fr")
    print("processing file: ", input_file)
    translated_doc = docx_translator.process_document()
    print("file processed: ", input_file)
    translated_file = BytesIO()
    translated_doc.save(translated_file)
    translated_file.seek(0)
    print("file saved: ", output_file)
    with open(output_file, "wb") as f:
        f.write(translated_file.getvalue())
        